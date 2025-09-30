from pathlib import Path
from re import findall

import random
from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
from random import choice
from string import ascii_lowercase

from ..Models.Fileschame import *

from .Cell import Cell
from .Table import Table as CustomTable


from .File import File


class DocxFile(File):
    def __init__(self, content: bytes):
        buffer: BytesIO = BytesIO(content)
        self.__file: Document = Document(buffer)

    @property
    def file(self):
        return self.__file

    def clear_cell(self, cell):
        cell._element.clear_content()  # удаляет весь внутренний XML

        # Создаем новый параграф
        paragraph = cell.add_paragraph()
        return paragraph

    def __generate_key(self, id_table: int, row_index: int, cell_index: int) -> str:
        low_name = f"t{id_table}_r{row_index}_c{cell_index}"

        return "{{" + low_name + "}}"

    def get_all_parser_table_in_file(self, re: str) -> list[tuple[int, str, CustomTable]]:
        tables = self.__file.tables
        custom_tables = []
        for i, table in enumerate(tables):
            if self.is_table_schema(table, re):
                key_table = findall(re, self.__get_key_table(table))[0]
                self.delete_system_row(i)
                custom_table = self.__create_table(table)
                custom_tables.append((i, key_table, custom_table))

        return custom_tables

    def rename_cell(self, id_table: int, name: str, rename: str):
        table = self.__file.tables[id_table]
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                 if cell.text == name:
                     cell.text = rename

    def generate_key_empty_cell(self, id_table: int):
        table = self.__file.tables[id_table]
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if not cell.text.strip():
                    cell.text = ""
                    key = self.__generate_key(id_table, row_index, cell_index)

                    paragraph = self.clear_cell(cell)

                    # Настраиваем выравнивание
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                    # Добавляем текст
                    paragraph.add_run(key)
        return self.__create_table(table)

    def render(self, schemas: dict):
        pass

    def __create_table(self, table: Table) -> CustomTable:
        custom_table = CustomTable(len(table.rows), len(table.columns))
        cells = self.get_list_cells(table)
        for cell in list(cells):
            width = cell[2][0] - cell[1][0]
            height = cell[2][1] - cell[1][1]
            custom_table.add_cell(cell[1][0], cell[1][1], width, height, cell[0])
        return custom_table

    def is_table_schema(self, table: Table, re: str) -> bool:
        key_table = self.__get_key_table(table)
        if len(key_table) > 0:
            count = len(findall(re, key_table))
            return count > 0
        return False

    def get_list_cells(self, table: Table):
        cells = set()
        for y, row in enumerate(table.columns):
            for x, cell in enumerate(row.cells):
                cell_loc = self.__get_coord_cell(cell)
                cells.add((cell.text, *cell_loc))
        return cells

    def delete_system_row(self, id_table: int) -> Table:
        self.__file.tables[id_table]._tbl.remove(self.__file.tables[id_table].rows[0]._tr)
        return self.__file.tables[id_table]

    def __get_key_table(self, table: Table) -> str | None:
        key_str = "".join(self.__get_text_cell(table, 0, 0).split()).lower()
        return key_str

    def __get_coord_cell(self, cell):
        tc = cell._tc
        try:
            return ((tc.left, tc.top), (tc.right, tc.bottom))
        except ValueError:
            return ((tc.left, tc.top), (tc.right, self.__get__bottom_coord_cell(cell._parent, cell, (tc.top, tc.left))))

    def __get__bottom_coord_cell(self, table: Table, cell, coord) -> int:
        text_cell = cell.text
        bottom_coord = coord[0]
        next_text_cell = table.cell(bottom_coord, coord[1])
        while text_cell == next_text_cell.text:
            bottom_coord += 1
            next_text_cell = table.cell(bottom_coord, coord[1])
        return bottom_coord

    def __create_list_cells_schames(self, cells: list[Cell], re: str):
        cells_schemas = [cell.get_schemas() for cell in cells]
        return cells_schemas

    def __get_text_cell(self, table: Table, x: int, y: int) -> str:
        return table.rows[x].cells[y].text
